"""
NeuralPath — Document Parser
Extracts plain text from PDF, DOCX, and TXT uploads.
"""
import io
from fastapi import UploadFile


async def extract_text(file: UploadFile) -> str:
    """Extract plain text from an uploaded PDF, DOCX, or TXT file."""
    content = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_docx(content)
    else:
        # Assume plain text (also handles .txt, .md, pasted text)
        return content.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF using pdfplumber (handles multi-column layouts)."""
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber not installed
        return _extract_pdf_fallback(data)
    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")


def _extract_pdf_fallback(data: bytes) -> str:
    """Fallback PDF extraction using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(
            page.extract_text() or ""
            for page in reader.pages
        ).strip()
    except Exception as e:
        raise ValueError(f"Could not read PDF (fallback also failed): {e}")


def _extract_docx(data: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {e}")
